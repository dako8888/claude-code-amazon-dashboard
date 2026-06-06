"""导出 Dashboard 使用手册为 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2); s.right_margin = Cm(2)

# Title
t = doc.add_heading("Amazon Workflow Dashboard 使用手册", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run("v1.0 | 2026-05-23"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
doc.add_paragraph()

# Quick start
doc.add_heading("快速启动", level=2)
for s in [
    "打开终端（Git Bash）",
    "cd E:\\WorkBuddy\\amazon-dashboard",
    "streamlit run app.py",
    "浏览器自动打开 http://localhost:8501",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph()

doc.add_heading("侧边栏", level=2)
doc.add_paragraph("左侧边栏有两个控件：")
for s in [
    "店铺选择：自有厨房 / 代运营家居 — 切换后所有页面数据即时切换",
    "功能菜单：8 个功能页面，点击切换",
    "底部显示当前数据目录和品类",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph()

# Page 1
doc.add_heading("1. 维护日历", level=2)
doc.add_paragraph("一屏看所有 ASIN 的健康状态和到期任务。")
doc.add_paragraph("显示内容：")
for s in [
    "每个 ASIN 一张卡片，含 BSR、评分、Search Terms 下次轮换日期",
    "红色边框 = 3 天内到期 | 黄色 = 7 天内 | 绿色 = 正常",
    "底部列出今天到期的事项",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("数据来源：/amazon-listing 创建 Listing 时自动写入的 state.json。")
doc.add_paragraph("使用频率：每天打开扫一眼。")
doc.add_paragraph()

# Page 2
doc.add_heading("2. ASIN 诊断", level=2)
doc.add_paragraph("深入看一个 ASIN 的当前指标 vs 基线，判断在涨还是在跌。")
doc.add_paragraph("功能：")
for s in [
    "选 ASIN → 左边看基线（创建日期/关键词数量），右边看维护信息（下次轮换/竞品检查时间）",
    "表单手动更新：CTR / CVR / BSR / 评分 / 近7天订单",
    "更新后自动重设维护提醒",
    "底部显示标题、5点内容、追踪的竞品列表",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("使用频率：每周一次，从亚马逊后台搬数据更新。")
doc.add_paragraph()

# Page 3
doc.add_heading("3. 竞品监控", level=2)
doc.add_paragraph("手动录入竞品快照，对比变化趋势。")
doc.add_paragraph("操作：")
for s in [
    "选 ASIN → 表单填竞品 ASIN / BSR / 价格 / 评分 / 评论数 / 标题 / 关键词",
    "保存后显示所有追踪竞品的最新快照",
    "多次录入同一竞品形成历史对比",
    "当前为手动模式 — MCP 激活后可自动拉取",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("使用频率：每周巡检竞品时录入。")
doc.add_paragraph()

# Page 4
doc.add_heading("4. 关键词管理", level=2)
doc.add_paragraph("看每个 ASIN 的关键词分布 + 执行 Search Terms 轮换。")
doc.add_paragraph("功能：")
for s in [
    "四列展示：核心大词 / 高转化长尾 / 场景属性 / Search Terms，各多少个",
    "轮换提醒：上次/下次轮换时间",
    "表单：填入新的 Search Terms → 自动校验 250 字节限制 → 执行",
    "注意：Search Terms 用空格分隔，不用逗号",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("使用频率：每 14-21 天轮换一次。")
doc.add_paragraph()

# Page 5
doc.add_heading("5. A/B 测试", level=2)
doc.add_paragraph("录入两个版本的数据，自动对比。")
doc.add_paragraph("操作：")
for s in [
    "选 ASIN → 填测试名称（如\"主图A/B测试\"）",
    "录入 A 版和 B 版的 CTR / CVR / 订单",
    "自动计算差异百分比 → 标出胜出版本",
    "注意：小样本波动不具有统计意义，仅供参考",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("使用频率：A/B 测试结束时录入。")
doc.add_paragraph()

# Page 6
doc.add_heading("6. 广告数据面板（v2.0 — 四象限智能分析）", level=2)
doc.add_paragraph("上传同一周期的所有广告报表，系统自动识别类型并交叉分析。不再简单粗暴地\"ACoS高=坏词\"。")
doc.add_paragraph("操作：")
for s in [
    "从亚马逊广告后台导出同一周期（如近7天）的所有报表 → .xlsx 格式",
    "支持类型: SP搜索词 / SP关键词 / SP Campaign / SB关键词 / SD投放 报告",
    "一次性多选上传 → 自动识别报表类型 → 交叉分析",
    "阈值基于你的实际数据动态计算，不写死数字",
]:
    p = doc.add_paragraph(s, style="List Bullet")

doc.add_paragraph("四象限分类：")
t4 = doc.add_table(rows=5, cols=3); t4.style = "Light Grid Accent 1"
for i, h in enumerate(["层级", "判定", "建议"]):
    t4.rows[0].cells[i].text = h
for ri, (label, judge, action) in enumerate([
    ("💰 金牛", "高花费 + 低ACoS", "拓词/加预算"),
    ("🔴 止损", "高花费 + ACoS>50%", "暂停或加否定关键词"),
    ("🟠 优化", "高花费 + ACoS偏高但可控", "降出价20-30%或换精准匹配"),
    ("🟡 潜力", "低花费 + 低ACoS", "加预算测试，关注规模增长"),
]):
    t4.rows[ri+1].cells[0].text = label
    t4.rows[ri+1].cells[1].text = judge
    t4.rows[ri+1].cells[2].text = action

doc.add_paragraph("")
doc.add_paragraph("MCP 自然排名交叉（Skill 1.0 联动）：")
for s in [
    "命令: python lib/ad_analyzer.py <报表目录> <店铺> --mcp <ASIN>",
    "自动查 SIF 自然排名 → 对比广告表现 → 发现浪费或机会",
    "例如: 某词自然排名第1但广告花费$50 → 建议降广告预算",
    "例如: 某词自然没有排名但广告转化好 → SEO优化机会",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("使用频率：每 8 天一次（SP 归因 7 天 + 1 天缓冲），或每 15 天（SB/SD 归因 14 天）。")
doc.add_paragraph()

# Page 7
doc.add_heading("7. 图片库存", level=2)
doc.add_paragraph("看每个 ASIN 的图片生成进度。")
doc.add_paragraph("显示内容：")
for s in [
    "5 个数字：总数 / 主图 / 辅图 / A+图 / AI生成",
    "进度条：主图 1/1 | 辅图 8/8 | A+ 5/5",
    "展开每类看具体文件名",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("使用频率：图片轨跑完后检查槽位是否缺图。")
doc.add_paragraph()

# Page 8
doc.add_heading("8. 季节性运营日历", level=2)
doc.add_paragraph("未来 90 天节日提醒 + 品类策略建议。无需任何数据，开箱即用。")
doc.add_paragraph("显示内容：")
for s in [
    "未来 90 天亚马逊节点（Prime Day / 黑五 / 圣诞 / 情人节 / 母亲节等）",
    "每个节点的准备开始日期",
    "进入准备期的标\"现在开始准备\"",
    "底部按品类给出季节性内容策略：",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph("厨房工具：Q4 圣诞饼干旺季 → Q1 情人节模具 → Q2 母亲节/毕业烘焙 → Q3 Prime Day")
doc.add_paragraph("家居装饰：Q4 圣诞新年最高峰 → Q1 情人节/春节 → Q2-Q3 婚礼季/花园派对")
doc.add_paragraph("使用频率：每月看一眼，大促前 30 天重点关注。")
doc.add_paragraph()

# Data sources
doc.add_heading("数据来源一览", level=2)
table = doc.add_table(rows=9, cols=3)
table.style = "Light Grid Accent 1"
for i, h in enumerate(["页面", "数据来源", "需要什么"]):
    table.rows[0].cells[i].text = h
rows = [
    ["维护日历", "state.json", "跑过 /amazon-listing"],
    ["ASIN 诊断", "state.json + 手动更新", "跑过 + 每周录入指标"],
    ["竞品监控", "手动录入 → state.json", "每周录入快照"],
    ["关键词管理", "state.json", "跑过 /amazon-listing"],
    ["A/B 测试", "手动录入", "AB结束录入"],
    ["广告数据面板", ".xlsx 上传", "亚马逊后台导出报表"],
    ["图片库存", "磁盘文件", "图片轨跑完即可"],
    ["季节性日历", "内置数据", "无 — 开箱即用"],
]
for ri, rd in enumerate(rows):
    for ci, text in enumerate(rd):
        table.rows[ri+1].cells[ci].text = text

doc.add_paragraph()
doc.add_heading("常见问题", level=2)
doc.add_paragraph("Q: 为什么大部分页面是空的？")
doc.add_paragraph("A: Dashboard 的数据来自 state.json。需要先通过 /amazon-listing 创建 Listing，生成 state.json 后才有内容。广告数据面板和季节性日历不需要 state.json，立刻可用。")
doc.add_paragraph("Q: 店铺切换后数据没变？")
doc.add_paragraph("A: 不同店铺数据物理隔离。切换店铺会重新读取该店铺目录下的所有 ASIN。")
doc.add_paragraph("Q: Search Terms 轮换怎么操作？")
doc.add_paragraph("A: 关键词管理页面 → 选 ASIN → 在表单填新词 → 点执行。系统会自动校验 250 字节限制，超了会报错提示删减。")

output = Path(r"C:\Users\Administrator\Desktop\Dashboard使用手册.docx")
doc.save(str(output))
print(f"[OK] {output}")
