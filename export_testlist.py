"""导出测试清单为 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2); s.right_margin = Cm(2)

t = doc.add_heading("Skill 1.0 测试清单", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run("2026-05-23 | 48 项 | 预计 2 天"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
doc.add_paragraph()

# Prep
doc.add_heading("测试前准备", level=2)
for s in [
    "选一个真实 ASIN（从代运营家居店铺，铃兰花灯品类）",
    "确认 LinkFox Agent 已充值（MCP 5个数据源可调用）",
    "确认千问 API 可用（settings.json 已配置 QWEN_API_KEY）",
    "确认 DeepSeek API 可用",
    "确认 FFmpeg 可用（source ~/.bashrc && which ffmpeg）",
    "确认 HyperFrames 可用（npx hyperframes --help）",
    "打开 Chrome + OpenCLI 扩展（daemon port 19825）",
    "清理旧的铃兰花灯测试数据（删除 B0B4HBGS9G 目录）",
]:
    p = doc.add_paragraph(s, style="List Bullet")
    for run in p.runs: run.font.size = Pt(10)
doc.add_paragraph()

# ============
SECTIONS = [
    ("Phase 0 — 策略分析", [
        ("P0-1", "MCP SIF 关键词反查", "用真实 ASIN 调 linkfox-sif-asin-keywords", "返回关键词+排名+搜索量"),
        ("P0-2", "MCP ABA 数据", "用一个关键词调 ABA 品牌分析", "返回搜索趋势/点击集中度"),
        ("P0-3", "MCP JungleScout 拓展", "用一个种子词调关键词拓展", "返回相关词+PPC竞价"),
        ("P0-4", "MCP 关键词竞争度", "调 SIF 关键词概览", "返回供需比/竞品数量"),
        ("P0-5", "MCP 关键词流量", "调 SIF 关键词流量拆解", "返回竞品流量来源"),
        ("P0-6", "ASIN 自诊断", "跑 self_diagnosis.py", "7项检查+三色分级"),
        ("P0-7", "竞品分级 T1/T2/T3", "指定 3-5 竞品 ASIN → AI 分级", "分级合理"),
        ("P0-8", "评论挖掘", "MCP拉竞品评论 → AI提取痛点", "输出结构化洞察"),
        ("P0-9", "COSMO 意图图谱", "DeepSeek 按6种类型归类关键词", "覆盖 ≥4 种关系类型"),
        ("P0-10", "策略简报输出", "汇总以上 → 输出 .docx", "含卖点排序+差异化+视觉方向"),
    ]),
    ("Phase 1 — 文字轨", [
        ("T1-1", "DeepSeek 标题生成", "输入Phase 0策略 → 生成标题", "≤200字符，公式覆盖≥3/5"),
        ("T1-2", "DeepSeek 5点生成", "每条对应 COSMO 意图", "5条，每条≤500字符"),
        ("T1-3", "DeepSeek 产品描述", "HTML 格式", "≤2000字符"),
        ("T1-4", "DeepSeek Search Terms", "自动去重补漏", "≤250字节，无逗号"),
        ("T1-5", "DeepSeek Alexa Q&A", "5组语音问答", "覆盖价格/功能/尺寸/场景"),
        ("T1-6", "文案 .docx 输出", "一键输出完整文案包", "标题+5点+描述+ST+QA"),
        ("T1-7", "文字质检", "python qa_text.py <ASIN> <store>", "PASS（0 hard_fail）"),
    ]),
    ("Phase 1 — 图片轨", [
        ("I1-1", "A+ 页面架构师", "aplus_designer.py 自动推荐模块", "匹配合适的产品原型"),
        ("I1-2", "A+ 提示词手写", "AI 读 Phase 0 → 写英文提示词", "中文目标+英文提示词+尺寸"),
        ("I1-3", "A+ 提示词 .docx", "export_prompts_docx() 导出", "格式正确"),
        ("I1-4", "Listing 9槽位提示词", "主图+8辅图全部生成", "每个槽位有独立英文提示词"),
        ("I1-5", "LinkFox 商品套图", "打开 LinkFox → 贴提示词 → Image2", "出图可用"),
        ("I1-6", "图片下载归档", "下载原图 → images/aplus/", "按模块分类"),
        ("I1-7", "图片质检", "python qa_images.py <ASIN> <store>", "尺寸+千问评价 PASS"),
    ]),
    ("Phase 1 — 视频轨", [
        ("V1-1", "视频策略简报", "从模板填 → video_strategy.md", "决策类型/路径/调性"),
        ("V1-2", "主图视频分镜", "选 main_video_35s.md 模板填", "8镜完整分镜"),
        ("V1-3", "分镜 .docx 导出", "export_storyboard_docx()", "中文目标+英文+参数"),
        ("V1-4", "Agent 指令生成", "填 agent_instruction_template.md", "英文指令完整"),
        ("V1-5", "HyperFrames 路径A", "写 index.html → npm run render", "30s MP4 产品视频"),
        ("V1-6", "LibTV 路径B", "打开画布 → 贴Agent指令 → 等出片", "30s 产品视频"),
        ("V1-7", "视频质检", "python qa_video.py <ASIN> <store> --qwen", "元数据+千问 PASS"),
    ]),
    ("Phase 2 — 质检汇总", [
        ("Q-1", "文字质检完整", "验证侵权/品牌/合规所有检查项", "无漏检"),
        ("Q-2", "图片质检完整", "尺寸+合规+千问三层都跑", "无漏检"),
        ("Q-3", "视频质检完整", "7种视频类型规格都验证", "无漏检"),
        ("Q-4", "质检报告保存", "JSON 保存到 qa/ 目录", "路径正确"),
    ]),
    ("Phase 3 — 维护", [
        ("M-1", "state.json 生成", "/amazon-listing 跑完自动写入", "video/maintenance/keywords 字段完整"),
        ("M-2", "版本管理", "重复跑同一 ASIN → 自动归档", "archive/ 有历史版本"),
        ("M-3", "Search Terms 轮换", "检查 state.json 到期日", "14天后提醒"),
    ]),
    ("Dashboard", [
        ("D-1", "启动", "streamlit run app.py → 浏览器打开", "无报错"),
        ("D-2", "广告数据面板", "上传报表 → 四象限分析", "分级合理"),
        ("D-3", "图片库存", "选 ASIN → 看槽位进度", "显示主图/辅图/A+完成度"),
        ("D-4", "维护日历", "有 state.json 的 ASIN", "卡片+到期标红"),
        ("D-5", "竞品监控", "手动录入 1 个竞品快照", "保存+显示"),
        ("D-6", "季节性日历", "查看未来 90 天节点", "显示准备提醒"),
    ]),
    ("广告分析引擎 v2.0", [
        ("A-1", "单报告分析", "python lib/ad_analyzer.py <目录>", "四象限+动态阈值"),
        ("A-2", "多报告交叉", "同周期多种报告一起上传", "自动识别+合并分析"),
        ("A-3", "MCP 自然排名交叉", "--mcp <ASIN>", "自然排名 vs 广告表现"),
        ("A-4", "docx 导出", "自动生成报告", "格式正确"),
    ]),
]

for title, items in SECTIONS:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=len(items)+1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["编号", "测试项", "怎么测", "预期结果"]):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(9); run.font.bold = True
    for ri, (nid, name, how, expect) in enumerate(items):
        for ci, text in enumerate([nid, name, how, expect]):
            cell = table.rows[ri+1].cells[ci]; cell.text = text
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(9)
    doc.add_paragraph()

# Test order
doc.add_heading("测试顺序建议", level=2)
for s in [
    "Day 1 上午: Phase 0（MCP验证+策略分析）→ Phase 1 文字（DeepSeek+质检）",
    "Day 1 下午: Phase 1 图片（LinkFox提示词+生图+质检）→ Dashboard 启动验证",
    "Day 2: Phase 1 视频（HyperFrames + LibTV 出片）→ 广告分析 → Phase 3 维护 → 端到端验证",
]:
    p = doc.add_paragraph(s, style="List Bullet")
doc.add_paragraph()

# Success criteria
doc.add_heading("成功标准", level=2)
for s in [
    "1 个真实 ASIN 完成 Phase 0 → Phase 3 全链路",
    "3 条 QA 命令（文字/图片/视频）全部 PASS（0 hard_fail）",
    "至少 1 个 MP4 视频产出",
    "Dashboard 广告面板返回合理分析",
    "state.json 字段完整",
    "再拿第 2 个 ASIN 全流程走一遍，无需改代码即可复用",
]:
    p = doc.add_paragraph(s, style="List Bullet")

output = Path(r"C:\Users\Administrator\Desktop\Skill1.0测试清单.docx")
doc.save(str(output))
print(f"[OK] {output}")
